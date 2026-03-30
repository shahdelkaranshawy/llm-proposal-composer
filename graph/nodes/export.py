from typing import Dict, Any, List, Tuple
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from graph.state import GraphState
import copy
import re


def parse_and_add_table(doc: Document, table_text: str) -> any:
    """
    Parse table text in markdown format and add as a Word table.
    Format: [TABLE]\n| Col1 | Col2 |\n|------|------|\n| Data1 | Data2 |\n[/TABLE]
    """
    try:
        # Extract table content between [TABLE] and [/TABLE]
        table_match = re.search(r'\[TABLE\](.*?)\[/TABLE\]', table_text, re.DOTALL)
        if not table_match:
            return None
        
        table_content = table_match.group(1).strip()
        
        # Split into rows
        rows = [r.strip() for r in table_content.split('\n') if r.strip() and r.strip().startswith('|')]
        
        if len(rows) < 2:  # Need at least header and one data row
            return None
        
        # Parse header row
        header_row = [cell.strip() for cell in rows[0].split('|')[1:-1]]  # Remove first and last empty splits
        
        # Skip separator row (the one with ----)
        data_rows = []
        for row in rows[1:]:
            if '---' not in row:  # Skip separator rows
                cells = [cell.strip() for cell in row.split('|')[1:-1]]
                if cells:  # Only add non-empty rows
                    data_rows.append(cells)
        
        if not data_rows:
            return None
        
        # Create Word table
        num_cols = len(header_row)
        num_rows = len(data_rows) + 1  # +1 for header
        
        table = doc.add_table(rows=num_rows, cols=num_cols)
        
        # Apply table style with better error handling
        try:
            # Try common table styles
            for style_name in ['Table Grid', 'Light List', 'Medium Grid 1', 'Light Grid']:
                try:
                    table.style = style_name
                    break
                except:
                    continue
        except Exception as e:
            print(f"      Note: Using default table style (no styling available)")
        
        # Add table borders manually if no style worked
        table.autofit = True
        
        # Add header row
        header_cells = table.rows[0].cells
        for i, header_text in enumerate(header_row):
            if i < len(header_cells):
                header_cells[i].text = header_text
                # Make header bold
                try:
                    for paragraph in header_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                        # If no runs, make the first run bold
                        if not paragraph.runs:
                            run = paragraph.add_run()
                            run.bold = True
                except:
                    pass
        
        # Add data rows
        for row_idx, row_data in enumerate(data_rows, start=1):
            if row_idx < len(table.rows):
                cells = table.rows[row_idx].cells
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx < len(cells):
                        cells[col_idx].text = cell_text
        
        return table
        
    except Exception as e:
        print(f"   ⚠️ Error parsing table: {e}")
        return None


def find_section_in_document(doc: Document, section_title: str) -> Tuple[int, int]:
    """
    Find a section in the document by title.
    Returns (start_index, end_index) of paragraphs, or (-1, -1) if not found.
    """
    start_idx = -1
    end_idx = -1
    
    for i, para in enumerate(doc.paragraphs):
        # Check if this paragraph is a heading with the section title
        if para.style.name.startswith('Heading') or para.text.strip().lower() == section_title.lower():
            if section_title.lower() in para.text.lower():
                start_idx = i
                # Find the end (next heading or end of document)
                for j in range(i + 1, len(doc.paragraphs)):
                    if doc.paragraphs[j].style.name.startswith('Heading'):
                        end_idx = j - 1
                        break
                if end_idx == -1:
                    end_idx = len(doc.paragraphs) - 1
                break
    
    return (start_idx, end_idx)


def section_has_table(doc: Document, start_idx: int, end_idx: int) -> bool:
    """
    Check if a section in the document contains a table.
    """
    try:
        # Get the parent element of the section
        if start_idx < 0 or start_idx >= len(doc.paragraphs):
            return False
        
        section_element = doc.paragraphs[start_idx]._element
        parent = section_element.getparent()
        
        # Look for table elements between start and end
        heading_index = parent.index(section_element)
        
        # Check elements after the heading for tables
        for i in range(heading_index + 1, len(parent)):
            element = parent[i]
            # Check if it's a table element
            if element.tag.endswith('tbl'):  # Word tables have 'tbl' tag
                # Check if this table is before the next section
                return True
            # If we hit another heading, stop looking
            if element.tag.endswith('p'):
                try:
                    para_idx = [p._element for p in doc.paragraphs].index(element)
                    if para_idx > end_idx:
                        break
                    para = doc.paragraphs[para_idx]
                    if para.style.name.startswith('Heading'):
                        break
                except:
                    pass
        
        return False
    except:
        return False


def replace_section_content(doc: Document, section_title: str, new_content: str) -> bool:
    """
    Replace the content of a specific section in the document.
    Preserves the heading and only replaces the content paragraphs.
    SKIPS sections that should preserve template content (like Table of Contents) or have existing tables.
    Returns True if section was found and replaced, False otherwise.
    """
    # List of sections to NEVER modify (preserve template content completely)
    preserve_sections = [
        'table of contents',
        'toc',
        'contents',
        'copyright',
        'document information',
        'document change record',
        'document reviewers',
        'prepared for',
        'change record',
        'reviewers'
    ]
    
    # Normalize section title for comparison (remove extra spaces, lowercase)
    normalized_title = ' '.join(section_title.lower().split())
    
    # Check if this section should be preserved
    if any(preserve_term in normalized_title for preserve_term in preserve_sections):
        print(f"   ⊙ Preserving template content for: {section_title} (administrative section with tables)")
        return False  # Don't replace, keep template as-is
    
    start_idx, end_idx = find_section_in_document(doc, section_title)
    
    if start_idx == -1:
        print(f"   ⚠️  Section '{section_title}' not found in template")
        return False
    
    # Check if this section already has a table in the template
    if section_has_table(doc, start_idx, end_idx):
        print(f"   ⊙ Preserving template content for: {section_title} (contains existing table)")
        return False  # Don't replace sections with existing tables
    
    print(f"   ✓ Found section '{section_title}' at paragraphs {start_idx}-{end_idx}")
    
    # Delete old content paragraphs (keep the heading)
    paragraphs_to_delete = []
    for i in range(start_idx + 1, end_idx + 1):
        if i < len(doc.paragraphs):
            paragraphs_to_delete.append(i)
    
    # Delete in reverse order to maintain indices
    for i in reversed(paragraphs_to_delete):
        p = doc.paragraphs[i]._element
        p.getparent().remove(p)
    
    # Insert new content after the heading
    heading_element = doc.paragraphs[start_idx]._element
    parent = heading_element.getparent()
    heading_index = parent.index(heading_element)
    
    # Process content - split by tables first, then by paragraphs
    parts = re.split(r'(\[TABLE\].*?\[/TABLE\])', new_content, flags=re.DOTALL)
    
    print(f"      Processing content: found {len(parts)} parts")
    tables_found = sum(1 for p in parts if '[TABLE]' in p)
    print(f"      Tables in content: {tables_found}")
    
    insert_position = heading_index + 1
    
    for part in parts:
        if not part.strip():
            continue
            
        # Check if this part is a table
        if part.strip().startswith('[TABLE]'):
            print(f"      Creating Word table...")
            # Create and insert table
            table = parse_and_add_table(doc, part)
            if table:
                print(f"      ✓ Table created: {len(table.rows)} rows × {len(table.columns)} cols")
                # Move table to correct position
                table_element = table._element
                parent.remove(table_element)
                parent.insert(insert_position, table_element)
                insert_position += 1
            else:
                print(f"      ✗ Table parsing failed")
        else:
            # Process as regular paragraphs
            content_paragraphs = part.split('\n\n')
            
            for para_text in content_paragraphs:
                if para_text.strip():
                    # Check if this is a subsection heading (starts with ##)
                    if para_text.strip().startswith('### '):
                        subsection_title = para_text.strip()[4:].strip()  # Remove ###
                        new_para = doc.add_heading(subsection_title, level=3)
                    elif para_text.strip().startswith('## '):
                        subsection_title = para_text.strip()[3:].strip()  # Remove ##
                        new_para = doc.add_heading(subsection_title, level=2)
                    else:
                        # Regular paragraph
                        new_para = doc.add_paragraph(para_text.strip())
                    
                    # Move to the right position
                    new_para_element = new_para._element
                    parent.remove(new_para_element)
                    parent.insert(insert_position, new_para_element)
                    insert_position += 1
    
    print(f"   ✓ Replaced content for '{section_title}'")
    return True


def export_node(state: GraphState) -> Dict[str, Any]:
    """
    Export the final proposal as a Word document using the template as a base.
    This preserves ALL template formatting, images, headers, footers, etc.
    """
    concatenated_sections = state.get("concatenated_sections", [])
    proposal_outline = state.get("proposal_outline", {})
    
    if not concatenated_sections:
        print("⚠️ No sections to export")
        return {"exported_document": None, "current_step": "export"}
    
    # Get the template file
    template_file = state.get("template_file")
    
    if not template_file:
        print("⚠️ No template file available - creating simple document")
        # Fallback: create a simple document
        doc = Document()
        for section in concatenated_sections:
            doc.add_heading(section.get("section_title", "Untitled"), level=1)
            doc.add_paragraph(section.get("content", ""))
    else:
        # Load the template as the base document
        try:
            template_file.seek(0)
            doc = Document(template_file)
            print(f"✅ Loaded template document ({len(doc.paragraphs)} paragraphs)")
            print(f"   Preserving: headers, footers, images, logos, formatting, page layout")
            
            # Get the proposal outline to know which sections use template content
            all_sections = proposal_outline.get("sections", [])
            section_flags = {}
            for s in all_sections:
                title = s.get("title", s.get("section_title", ""))
                use_template = s.get("use_template_content", False)
                section_flags[title] = use_template
            
            print(f"\n📝 Processing {len(concatenated_sections)} sections...")
            print(f"   Sections marked to use template content: {[k for k, v in section_flags.items() if v]}")
            
            sections_replaced = 0
            sections_appended = []
            
            # Define sections that should always be preserved from template
            preserve_section_keywords = [
                'table of contents', 'toc', 'contents',
                'copyright', 'document information',
                'document change record', 'document reviewers',
                'prepared for', 'change record', 'reviewers'
            ]
            
            # Try to find and replace each generated section in the template
            for section in concatenated_sections:
                section_title = section.get("section_title", section.get("title", ""))
                section_content = section.get("content", "")
                use_template = section_flags.get(section_title, False)
                
                # Normalize section title for comparison
                normalized_title = ' '.join(section_title.lower().split())
                
                # Check if this is a section that should be preserved from template
                should_preserve = any(keyword in normalized_title for keyword in preserve_section_keywords)
                
                # CRITICAL: If use_template_content=true OR it's an administrative section, ALWAYS preserve
                if use_template or should_preserve:
                    print(f"   ⊙ PRESERVING template content for: {section_title} (use_template={use_template})")
                    # Don't process this section at all - leave template exactly as-is
                    continue
                
                # For sections that need new content, try to find and replace in template
                print(f"   → Replacing content for: {section_title}")
                replaced = replace_section_content(doc, section_title, section_content)
                
                if replaced:
                    sections_replaced += 1
                else:
                    # Section not found in template - will append at end
                    print(f"      Section not in template, will append")
                    sections_appended.append(section)
            
            print(f"\n✅ Replaced {sections_replaced} sections in template")
            
            # Append any sections that weren't found in template
            if sections_appended:
                print(f"   Adding {len(sections_appended)} new sections at the end...")
                doc.add_page_break()
                
                for section in sections_appended:
                    section_title = section.get("section_title", section.get("title", "Untitled"))
                    section_content = section.get("content", "")
                    
                    doc.add_heading(section_title, level=1)
                    
                    if section_content:
                        # Split by tables first
                        parts = re.split(r'(\[TABLE\].*?\[/TABLE\])', section_content, flags=re.DOTALL)
                        
                        for part in parts:
                            if not part.strip():
                                continue
                            
                            # Check if this is a table
                            if part.strip().startswith('[TABLE]'):
                                table = parse_and_add_table(doc, part)
                                if table:
                                    print(f"      ✓ Added table to {section_title}")
                            else:
                                # Process paragraphs
                                paragraphs = part.split('\n\n')
                                for para in paragraphs:
                                    if para.strip():
                                        # Handle subsection headings
                                        if para.strip().startswith('### '):
                                            doc.add_heading(para.strip()[4:].strip(), level=3)
                                        elif para.strip().startswith('## '):
                                            doc.add_heading(para.strip()[3:].strip(), level=2)
                                        else:
                                            doc.add_paragraph(para.strip())
                    
                    doc.add_paragraph()
                
                print(f"   ✓ Added {len(sections_appended)} new sections")
            
            print(f"\n🎉 Document ready: Template formatting preserved, content updated!")
            
        except Exception as e:
            print(f"❌ Error processing template: {e}")
            import traceback
            traceback.print_exc()
            # Fallback to simple document
            doc = Document()
            for section in concatenated_sections:
                doc.add_heading(section.get("section_title", "Untitled"), level=1)
                doc.add_paragraph(section.get("content", ""))
    
    # Save document to BytesIO
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    return {
        "exported_document": doc_buffer,
        "current_step": "export"
    }
